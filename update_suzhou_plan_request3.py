from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path(r"C:\Users\aw4wz\Documents\Codex\TLaser\SuzhouIndustrialPark_TLaser_filled.docx")


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_text(p, text, size=10.5):
    for r in list(p.runs):
        p._p.remove(r._r)
    for i, part in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run_font(run, size=size)
    p.paragraph_format.line_spacing = 1.15


def set_cell_text(cell, text, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_text(p, text, size=size)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), "100")
        node.set(qn("w:type"), "dxa")


def replace_text_in_runs(doc, old, new):
    for p in doc.paragraphs:
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)


def replace_para_start(doc, prefix, text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            set_paragraph_text(p, text)
            return True
    return False


doc = Document(str(DOCX))

# 1. Correct customer names from scanned-contract context.
for old, new in [("北京海科", "北京海基"), ("上海盈耀", "上海应尧")]:
    replace_text_in_runs(doc, old, new)

# 2. Update customer targeting with Suzhou companies first.
summary = (
    "TLaser 是面向光通讯半导体二极管激光器的实时数字孪生与参数标定平台，目标是把传统依赖专家经验、离线仿真和反复流片/测试的激光器设计流程，升级为可在线校准、可快速预测、可持续积累数据的 AI+工业软件平台。项目以前期 CAE 软件算法、硅光仿真核心算法、多物理场仿真平台 MySim、高频电磁仿真 EML 以及数字健康仿真 HemoDyn 的开发经验为基础。申报人负责完成的项目合同附件显示，其曾围绕 EML 开发、EastWave 电磁仿真、Ansys 中国、清华大学深圳研究生院、天津科技大学、上海市气象局信息中心、苏州 ShonCloud、北京海基、上海应尧等客户或机构开展软件算法与工程项目交付，具备把复杂数值算法转化为客户可用软件的实证基础。项目已完成 TLaser 原型系统：包括准三维高保真物理模拟器、7 维参数空间数据生成、物理信息神经网络（PINN）代理模型、L-I-V 曲线在线反演标定引擎和中英文交互式控制面板。现有流水线已通过数据生成、模型训练和参数标定的自动化冒烟测试，代理模型可将几何、温度、电流和镜面反射率等输入映射到光功率、效率以及载流子/光子空间剖面，单次预测目标低于 5 毫秒。"
    "\n项目团队核心优势在于“计算物理 + 工业软件 + 人工智能 + 微波/高速电子系统”的复合能力。申报人万振文自 1996 年起长期从事海洋、生态、多物理场和 CAE 算法研究，2016 年后专注 CAE 软件与工程算法开发，并创办 MySim Digital Technology ApS；其作为发明人的已授权发明专利包括“一种有限元分析系统、方法、计算机设备及存储介质”（ZL 2019 1 1067815.0）、“海底挖沟工程预测方法、装置及服务器”（ZL 2019 1 0530513.6）和“绿潮生物量预报方法、装置、设备及介质”（ZL 2019 1 1249340.7），证明其在数值算法、预测模型和软件化知识产权方面有持续产出。拟加入团队成员秦三团为西安电子科技大学博士，具备 3G/4G/5G 移动通信系统、射频电路、数字硬件、FPGA/ARM 嵌入式软件和微波通信系统项目研发经验，可补强 TLaser 面向光模块客户的高速电子、测试系统和工程交付能力。团队拟联合半导体物理、电磁仿真和光电子工艺专家，形成从物理建模、数值求解、机器学习、云端产品到行业交付的闭环。"
    "\n目标产品为“光通讯激光器设计自动化和性能标定云端服务平台”。优先服务苏州工业园区及苏州光通信产业链企业，包括旭创科技/中际旭创、度亘核芯、苏纳光电、苏州易缆微、东辉光学、奇点光子、荣旗科技等光模块、激光芯片、光学元件、硅光芯片和智能检测装备企业；再逐步拓展至长三角和全国半导体激光器设计公司、光模块企业、晶圆代工/封测企业和科研院所，提供设计参数筛选、批次一致性分析、器件漂移诊断、良率提升和数字化研发服务。未来三年项目将以苏州工业园区为产业化基地，第一年完成公司注册、核心团队落地、原型系统工程化和首批试点客户验证；第二年完成多器件型号适配、云端 SaaS 版本和软著/专利布局，形成付费试点；第三年完成商业化版本、行业案例和规模化销售。"
)
replace_para_start(doc, "TLaser 是面向光通讯半导体二极管激光器", summary)

t3 = doc.tables[3]
set_cell_text(t3.cell(3, 0), "万振文：启明计划人才，计算物理与 CAE 算法专家，具备从科研建模、工业软件、数值算法、合同交付到创业产品化的连续经验。公开网页资料显示，其自 2016 年以来聚焦 CAE 软件算法开发，曾独立开发紧凑二维几何库和表面网格库，并创办 MySim Digital Technology ApS。附件专利证书显示，其作为发明人已取得有限元分析系统、海底挖沟工程预测和绿潮生物量预报等发明专利授权；附件合同显示其在高校、科研机构、工业软件企业和行业客户中完成过多类算法项目交付。秦三团：西安电子科技大学博士，公开资料显示其曾在中兴通讯从事 3G/4G/5G 移动通信系统研发，历任电路硬件研发工程师、整机系统研发工程师、项目技术总监；2014 年后在西安邮电大学创立微波通信实验室及研发团队，研究方向覆盖射频电路、数字硬件、FPGA/ARM 嵌入式软件和微波通信系统，可补强项目在光模块高速电子接口、测试标定硬件、微波/射频测量和工程化交付方面的能力。团队优势在于可把复杂物理问题转化为工程软件产品，并能与苏州园区光电子、集成电路、人工智能和工业软件生态形成协同。")
set_cell_text(t3.cell(14, 0), "营销策略：优先采用“苏州园区标杆客户联合验证 + 订阅服务 + 定制算法模块”的进入策略。第一阶段将苏州工业园区和苏州本地光通信企业放在首位，重点面向旭创科技/中际旭创、度亘核芯、苏纳光电、苏州易缆微、东辉光学、奇点光子、荣旗科技等光模块、激光芯片、光学元件、硅光芯片和智能检测装备企业，围绕参数筛选、L-I-V 标定、批次漂移诊断和测试数据闭环建立试点；第二阶段按项目制收取模型适配费、数据清洗费和算法服务费；第三阶段形成云端 SaaS 订阅、私有化部署和企业年度维护服务。定价以节省的仿真时间、测试成本、研发人员工时和良率提升收益为依据，建立苏州标杆案例后复制到长三角和全国光电子器件产业链。")

replace_para_start(doc, "TLaser 项目负责人万振文具备计算物理", "TLaser 项目负责人万振文具备计算物理、CAE 软件、优化算法和工程化原型开发的复合背景。公开网页资料显示，申报人自 2016 年以来聚焦 CAE 软件算法开发与咨询，擅长 FEM、FVM、FDM 和优化算法，曾独立开发二维几何库和表面网格库；此前 1996-2016 年长期从事海洋与生态建模研究，拥有丹麦气象研究所、厦门大学等科研经历。申报人负责完成的合同附件覆盖 EML 开发、EastWave 电磁仿真、Ansys 中国、高校和行业客户软件算法项目，显示其具备跨行业算法软件交付能力。秦三团拟作为核心团队成员之一加入项目，其为西安电子科技大学博士，公开资料显示曾在中兴通讯参与或主持 3G/4G/5G 无线通信系统开发，2014 年后创立微波通信实验室及师生研发团队，承担数十项军工及民用微波通信相关项目，研究方向包括射频电路、数字硬件、FPGA/ARM 嵌入式软件。该背景与光模块测试、激光器驱动、L-I-V 测量和高速电子接口具有较强互补性。项目拟联合半导体物理、电磁场仿真、光电子器件测试和云端软件工程人才，建设面向光通讯激光器研发与量产测试的数字孪生平台。\n\n行业背景方面，光通讯正向 800G/1.6T、数据中心互联、硅光集成和国产光模块供应链升级发展，半导体激光器作为核心光源，对效率、温漂、可靠性和一致性要求持续提高。苏州工业园区已形成以高端光模块为牵引、核心光学元器件为支撑、光芯片与智能检测装备为突破的光通信产业生态，TLaser 将优先服务园区和苏州客户，包括旭创科技/中际旭创、度亘核芯、苏纳光电、苏州易缆微、东辉光学、奇点光子、荣旗科技等，再拓展至长三角和全国光通信、硅光、激光芯片与测试装备企业。")

replace_para_start(doc, "项目定位于光电子工业软件和 AI+半导体器件研发工具", "项目定位于光电子工业软件和 AI+半导体器件研发工具，重点服务光通讯半导体激光器的设计、测试和量产一致性分析。市场进入顺序上，TLaser 将把苏州工业园区及苏州客户放在最前面：一是面向旭创科技/中际旭创等高端光模块企业，提供光源模型、L-I-V 数据闭环和供应链器件一致性分析；二是面向度亘核芯等激光芯片/泵浦模块企业，提供器件参数反演、可靠性漂移诊断和设计参数筛选；三是面向苏纳光电、东辉光学、苏州易缆微、奇点光子等光学元件、硅光/光电芯片和光互连企业，提供器件级数字孪生、耦合与测试数据建模接口；四是面向荣旗科技等智能检测装备企业，探索 TLaser 与自动化测试平台的数据接口和算法嵌入。随着数据中心、AI 算力集群、5G/6G 承载网、硅光集成和高速光模块升级，通信级激光器对效率、带宽、可靠性和温度稳定性的要求持续提高；同时国内企业在高端 EDA/CAE、光电子器件仿真和数字孪生工具上仍存在国产替代空间。TLaser 不与通用有限元/电磁仿真软件正面竞争，而是聚焦“光通讯激光器 + 实测数据闭环 + 毫秒级代理模型”的垂直场景，作为研发部门和测试部门之间的数字化连接工具。")

replace_para_start(doc, "目标客户包括光通讯激光器芯片设计企业", "目标客户优先锁定苏州工业园区和苏州光通信产业链企业，包括旭创科技/中际旭创、度亘核芯、苏纳光电、苏州易缆微、东辉光学、奇点光子、荣旗科技等；随后拓展到长三角及全国光通讯激光器芯片设计企业、光模块企业、硅光和光电子器件研发团队、晶圆代工/封测企业、测试设备企业和高校科研院所。市场进入将采用三步策略：第一，选择苏州本地标杆客户开展联合验证，以器件参数筛选、L-I-V 曲线标定、批次漂移诊断等具体场景证明价值；第二，形成按项目收费的模型适配和数据标定服务，积累可复用器件模型库；第三，推出云端 SaaS 订阅、私有化部署和年度维护服务。定价策略以客户节省的仿真工时、测试样本、流片迭代和良率改善为依据，早期试点价格控制在客户易接受范围，商业版按账号数、器件型号数、数据量和私有化部署范围分级收费。")

# 3. Add Qin Santuan to the team member table.
t15 = doc.tables[15]
set_cell_text(t15.cell(0, 1), "秦三团")
set_cell_text(t15.cell(0, 5), "待补充")
set_cell_text(t15.cell(1, 1), "男")
set_cell_text(t15.cell(1, 5), "中国")
set_cell_text(t15.cell(2, 5), "博士；微波系统研发专家")
set_cell_text(t15.cell(3, 1), "博士")
set_cell_text(t15.cell(3, 5), "西安电子科技大学")
set_cell_text(t15.cell(4, 1), "工学")
set_cell_text(t15.cell(4, 5), "电子科学与技术/电磁场与微波技术")
set_cell_text(t15.cell(5, 3), "□身份证  □护照号（待补充）")
set_cell_text(t15.cell(5, 8), "待补充")
set_cell_text(t15.cell(6, 3), "□有  ☑无（公开资料未见海外经历）")
set_cell_text(t15.cell(6, 8), "西安")
set_cell_text(t15.cell(7, 3), "□是  ☑否（拟按项目需要来苏州）")
set_cell_text(t15.cell(7, 8), "待安排")
set_cell_text(t15.cell(8, 3), "□全职  ☑兼职（拟，技术顾问/联合研发）")
set_cell_text(t15.cell(8, 8), "比例待定")
set_cell_text(t15.cell(9, 3), "目前任职及科研团队在西安，拟以顾问或联合研发方式参与项目。")
set_cell_text(t15.cell(10, 3), "待补充")
set_cell_text(t15.cell(10, 8), "qinsantuan@126.com")
set_cell_text(t15.cell(13, 0), "2014.07")
set_cell_text(t15.cell(13, 2), "至今")
set_cell_text(t15.cell(13, 4), "西安邮电大学")
set_cell_text(t15.cell(13, 7), "微波通信实验室负责人、硕士生导师")
set_cell_text(t15.cell(14, 0), "2005.07")
set_cell_text(t15.cell(14, 2), "2014.07")
set_cell_text(t15.cell(14, 4), "中兴通讯股份有限公司")
set_cell_text(t15.cell(14, 7), "电路硬件/整机系统研发工程师、项目技术总监")
set_cell_text(t15.cell(20, 0), "博士阶段")
set_cell_text(t15.cell(20, 2), "毕业")
set_cell_text(t15.cell(20, 4), "西安电子科技大学")
set_cell_text(t15.cell(20, 7), "电子科学与技术")
set_cell_text(t15.cell(20, 9), "博士")
set_cell_text(t15.cell(26, 0), "就业经历和业绩：秦三团博士毕业于西安电子科技大学，公开资料显示其曾于 2005 年 7 月至 2014 年 7 月任职中兴通讯股份有限公司，从事 3G/4G/5G 移动通信系统研发，历任电路硬件研发工程师、整机系统研发工程师、项目技术总监，参与或主持多项无线通信系统开发。2014 年 7 月进入西安邮电大学后，创立微波通信实验室及师生研发团队，从事微波通信系统研究开发和人才培养，承担数十项军工及民用领域微波通信相关项目。其研究方向包括射频电路设计、数字硬件电路设计、FPGA/ARM 嵌入式软件设计，公开论文涉及瞬态电磁散射、TDIE/TDPO 混合算法、平台天线宽带分析等；公开专利包括“一种单音场强测量的方法与装置”和“一种人体静电测量装置”。\n\n与本项目关联：秦三团可承担 TLaser 面向光模块企业的测试硬件接口、激光器驱动与 L-I-V 采集、射频/高速电子系统、FPGA/嵌入式数据采集和客户现场工程化验证工作，与万振文的计算物理、数字孪生和 CAE 软件能力形成互补。")

doc.save(str(DOCX))
print(str(DOCX))
