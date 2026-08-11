from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table


DOCX = Path(r"C:\Users\aw4wz\Documents\Codex\TLaser\SuzhouIndustrialPark_TLaser_filled.docx")


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_text(p, text, size=10.5):
    for r in list(p.runs):
        p._p.remove(r._r)
    for i, part in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run_font(run, size=size)
    p.paragraph_format.line_spacing = 1.15


def set_cell_text(cell, text, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_text(p, text, size=size)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), "100")
        node.set(qn("w:type"), "dxa")


def replace_para_start(doc, prefix, text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            set_paragraph_text(p, text)
            return True
    return False


def table_text(table):
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def fill_jiang(table):
    set_cell_text(table.cell(0, 1), "蒋寻涯")
    set_cell_text(table.cell(0, 5), "待补充")
    set_cell_text(table.cell(1, 1), "男")
    set_cell_text(table.cell(1, 5), "中国")
    set_cell_text(table.cell(2, 5), "中科院百人计划；复旦大学研究员/博导")
    set_cell_text(table.cell(3, 1), "博士")
    set_cell_text(table.cell(3, 5), "Iowa State University")
    set_cell_text(table.cell(4, 1), "理学")
    set_cell_text(table.cell(4, 5), "物理学/光子学/电磁与光电仿真")
    set_cell_text(table.cell(5, 3), "□身份证  □护照号（待补充）")
    set_cell_text(table.cell(5, 8), "待补充")
    set_cell_text(table.cell(6, 3), "☑有  □无")
    set_cell_text(table.cell(6, 8), "上海")
    set_cell_text(table.cell(7, 3), "□是  ☑否（拟按项目需要来苏州）")
    set_cell_text(table.cell(7, 8), "待安排")
    set_cell_text(table.cell(8, 3), "□全职  ☑兼职（拟，技术顾问/联合研发）")
    set_cell_text(table.cell(8, 8), "比例待定")
    set_cell_text(table.cell(9, 3), "目前主要任职在复旦大学/相关光电仿真软件产业化平台，拟以顾问或联合研发方式参与项目。")
    set_cell_text(table.cell(10, 3), "待补充")
    set_cell_text(table.cell(10, 8), "jiangxunya@fudan.edu.cn")
    set_cell_text(table.cell(13, 0), "2023.06")
    set_cell_text(table.cell(13, 2), "至今")
    set_cell_text(table.cell(13, 4), "西湖大学光电研究院/上海东峻")
    set_cell_text(table.cell(13, 7), "光电芯片设计与验证中心主任/创始人")
    set_cell_text(table.cell(14, 0), "2012.10")
    set_cell_text(table.cell(14, 2), "至今")
    set_cell_text(table.cell(14, 4), "复旦大学")
    set_cell_text(table.cell(14, 7), "研究员、博士生导师")
    set_cell_text(table.cell(15, 0), "2003.03")
    set_cell_text(table.cell(15, 2), "2012.10")
    set_cell_text(table.cell(15, 4), "中科院上海微系统所")
    set_cell_text(table.cell(15, 7), "研究员、博士生导师")
    set_cell_text(table.cell(16, 0), "2002.09")
    set_cell_text(table.cell(16, 2), "2004.07")
    set_cell_text(table.cell(16, 4), "MIT 电子研究所")
    set_cell_text(table.cell(16, 7), "Research Associate")
    set_cell_text(table.cell(20, 0), "博士阶段")
    set_cell_text(table.cell(20, 2), "2001")
    set_cell_text(table.cell(20, 4), "Iowa State University")
    set_cell_text(table.cell(20, 7), "物理学")
    set_cell_text(table.cell(20, 9), "博士")
    set_cell_text(table.cell(21, 0), "1993.08")
    set_cell_text(table.cell(21, 2), "1996.07")
    set_cell_text(table.cell(21, 4), "北京大学")
    set_cell_text(table.cell(21, 7), "物理学")
    set_cell_text(table.cell(21, 9), "硕士")
    set_cell_text(table.cell(22, 0), "1985.08")
    set_cell_text(table.cell(22, 2), "1989.07")
    set_cell_text(table.cell(22, 4), "北京师范大学")
    set_cell_text(table.cell(22, 7), "物理学")
    set_cell_text(table.cell(22, 9), "本科")
    set_cell_text(table.cell(26, 0), "就业经历和业绩：蒋寻涯教授为复旦大学研究员、博士生导师，公开资料显示曾在美国 Ames 国家实验室、Dicon Fiberoptics（硅谷）、MIT 电子研究所、中科院上海微系统所、复旦大学等机构从事光子学、光电器件、电磁与光电仿真研究。其研究方向覆盖激光/LED、光子晶体、超材料、特异介质、非线性与无序体系等；公开资料显示其发表 SCI 论文逾百篇，被引用 3000 余次，出版 Springer 学术专著，拥有多项发明专利，并曾入选中科院百人计划、上海杰出引进人才、浦江计划、浙江省千人计划等人才项目。\n\n产业化业绩：蒋寻涯教授创办上海东峻信息科技有限公司，长期推进国产电磁/光电仿真软件 EastWave 产业化，公开报告资料显示该软件在天线阵、共形天线、频选、超材料、电磁兼容、RCS/成像、微波暗室等工程场景有应用，并销售给航空、航天、船舶、中电等企业和科研院所。与本项目关联：蒋教授可为 TLaser 提供光电仿真软件、半导体光电子器件、国产工业软件商业化和客户工程应用方面的指导。")


doc = Document(str(DOCX))

team_summary = (
    "万振文：启明计划人才，计算物理与 CAE 算法专家，具备从科研建模、工业软件、数值算法、合同交付到创业产品化的连续经验。公开网页资料显示，其自 2016 年以来聚焦 CAE 软件算法开发，曾独立开发紧凑二维几何库和表面网格库，并创办 MySim Digital Technology ApS。附件专利证书显示，其作为发明人已取得有限元分析系统、海底挖沟工程预测和绿潮生物量预报等发明专利授权；附件合同显示其在高校、科研机构、工业软件企业和行业客户中完成过多类算法项目交付。"
    "\n蒋寻涯：复旦大学研究员、博士生导师，曾入选中科院百人计划，长期从事光子学、电磁/光电仿真和国产工业软件产业化。公开资料显示，其曾在 Ames 国家实验室、硅谷光电公司、MIT、中科院上海微系统所、复旦大学等机构工作，研究方向覆盖激光/LED、光子晶体、超材料、非线性、无序体系等；创办上海东峻并推进 EastWave 国产电磁仿真软件商业化。蒋教授可补强 TLaser 在光电器件物理、全波仿真、工业软件产品化和高端客户应用方面的能力。"
    "\n秦三团：西安电子科技大学博士，公开资料显示其曾在中兴通讯从事 3G/4G/5G 移动通信系统研发，历任电路硬件研发工程师、整机系统研发工程师、项目技术总监；2014 年后在西安邮电大学创立微波通信实验室及研发团队，研究方向覆盖射频电路、数字硬件、FPGA/ARM 嵌入式软件和微波通信系统，可补强项目在光模块高速电子接口、测试标定硬件、微波/射频测量和工程化交付方面的能力。"
)
t3 = doc.tables[3]
set_cell_text(t3.cell(3, 0), team_summary)

replace_para_start(
    doc,
    "TLaser 项目负责人万振文具备计算物理",
    "TLaser 项目负责人万振文具备计算物理、CAE 软件、优化算法和工程化原型开发的复合背景。公开网页资料显示，申报人自 2016 年以来聚焦 CAE 软件算法开发与咨询，擅长 FEM、FVM、FDM 和优化算法，曾独立开发二维几何库和表面网格库；此前 1996-2016 年长期从事海洋与生态建模研究，拥有丹麦气象研究所、厦门大学等科研经历。蒋寻涯教授为复旦大学研究员、博士生导师，曾入选中科院百人计划，并在光子学、电磁/光电仿真软件、EastWave 产业化和高端工程客户应用方面有长期积累；秦三团博士毕业于西安电子科技大学，具备中兴通讯 3G/4G/5G 系统研发、射频/数字硬件、FPGA/ARM 和微波通信系统项目经验。三人的能力组合覆盖计算物理、光电子仿真、国产工业软件、高速电子与工程测试，适合建设面向光通讯激光器研发与量产测试的数字孪生平台。\n\n行业背景方面，光通讯正向 800G/1.6T、数据中心互联、硅光集成和国产光模块供应链升级发展，半导体激光器作为核心光源，对效率、温漂、可靠性和一致性要求持续提高。苏州工业园区已形成以高端光模块为牵引、核心光学元器件为支撑、光芯片与智能检测装备为突破的光通信产业生态，TLaser 将优先服务园区和苏州客户，包括旭创科技/中际旭创、度亘核芯、苏纳光电、苏州易缆微、东辉光学、奇点光子、荣旗科技等，再拓展至长三角和全国光通信、硅光、激光芯片与测试装备企业。"
)

# Existing table 15 currently contains Qin. Copy it once, fill original with Jiang,
# leave copied table as Qin so both members remain represented.
has_jiang_team_table = any(
    len(table.rows) >= 27 and "蒋寻涯" in table_text(table) and "EastWave" in table_text(table)
    for table in doc.tables
)
qin_table = None
for table in doc.tables:
    text = table_text(table)
    if "秦三团" in text and "中兴通讯" in text and len(table.rows) >= 27:
        qin_table = table
        break

if qin_table is not None and not has_jiang_team_table:
    qin_tbl_copy = deepcopy(qin_table._tbl)
    qin_table._tbl.addnext(qin_tbl_copy)
    fill_jiang(qin_table)

doc.save(str(DOCX))
print(str(DOCX))
