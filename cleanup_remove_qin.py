from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

DOCX = Path(r"C:\Users\aw4wz\Documents\Codex\TLaser\SuzhouIndustrialPark_TLaser_revised_0810_MaRong.docx")


def set_run_font(run, size=10.5):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)


def set_paragraph_text(p, text, size=10.5):
    for r in list(p.runs):
        p._p.remove(r._r)
    for i, part in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run_font(run, size)
    p.paragraph_format.line_spacing = 1.15


def set_cell_text(cell, text):
    cell.text = ""
    set_paragraph_text(cell.paragraphs[0], text)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), "100")
        node.set(qn("w:type"), "dxa")


doc = Document(str(DOCX))

# Replace the remaining stage-target text that came from the prior Qin version.
set_cell_text(
    doc.tables[3].cell(16, 0),
    "第一年：完成苏州公司注册，万振文和马戎为核心全职成员落地，蒋寻涯作为兼职首席科学顾问参与技术路线和测试验证评审；完成 TLaser V1.0 工程版，接入至少 1 家苏州客户真实 L-I-V 数据，主营收入 80 万元，新增就业 5 人。\n第二年：形成多器件型号模型库和 SaaS 试点版，完成 3 家以上付费试点，主营收入 350 万元，团队 12 人。\n第三年：形成商业版和私有化部署能力，完成 6 类以上器件型号验证，主营收入 900 万元，团队 25 人。三年累计受理发明专利 5 项、PCT 2 项、软著 6 项，力争高企培育/申报和园区/市级科技项目支持。"
)

jiang_text = (
    "就业经历和业绩：蒋寻涯教授为复旦大学研究员、博士生导师，公开资料显示曾在美国 Ames 国家实验室、Dicon Fiberoptics（硅谷）、MIT 电子研究所、中科院上海微系统所、复旦大学等机构从事光子学、光电器件、电磁与光电仿真研究。其研究方向覆盖激光/LED、光子晶体、超材料、特异介质、非线性与无序体系等；公开资料显示其发表 SCI 论文逾百篇，被引用 3000 余次，出版 Springer 学术专著，拥有多项发明专利，并曾入选中科院百人计划、上海杰出引进人才、浦江计划、浙江省千人计划等人才项目。\n\n"
    "产业化业绩：蒋寻涯教授创办上海东峻信息科技有限公司，长期推进国产电磁/光电仿真软件 EastWave 产业化。与本项目关联：蒋教授负责 TLaser 光电仿真模型、激光器 L-I-V 测试验证技术路线、测试数据与物理模型一致性评审，以及半导体光电子产业客户资源协同。"
)

for table in doc.tables:
    text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    if len(table.rows) >= 27 and "蒋寻涯" in text:
        set_cell_text(table.cell(26, 0), jiang_text)

doc.save(str(DOCX))
print(str(DOCX))
