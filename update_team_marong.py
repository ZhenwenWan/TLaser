from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path(r"C:\Users\aw4wz\Documents\Codex\TLaser\SuzhouIndustrialPark_TLaser_revised_0810.docx")
OUT = Path(r"C:\Users\aw4wz\Documents\Codex\TLaser\SuzhouIndustrialPark_TLaser_revised_0810_MaRong.docx")


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_text(p, text, size=10.5):
    for r in list(p.runs):
        p._p.remove(r._r)
    for i, part in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run_font(run, size=size)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)


def set_cell_text(cell, text, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_text(p, text, size=size)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), "100")
        node.set(qn("w:type"), "dxa")


def replace_para_start(doc, prefix, text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            set_paragraph_text(p, text)
            return True
    return False


def table_text(table):
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def fill_ma_rong(table):
    set_cell_text(table.cell(0, 1), "马戎")
    set_cell_text(table.cell(0, 5), "待补充")
    set_cell_text(table.cell(1, 1), "男")
    set_cell_text(table.cell(1, 5), "中国")
    set_cell_text(table.cell(2, 5), "深圳企业经营管理与商务拓展")
    set_cell_text(table.cell(3, 1), "本科")
    set_cell_text(table.cell(3, 5), "待补充")
    set_cell_text(table.cell(4, 1), "工学/管理")
    set_cell_text(table.cell(4, 5), "自动化装备、电子开发、客户商务")
    set_cell_text(table.cell(5, 3), "□身份证  □护照号（待补充）")
    set_cell_text(table.cell(5, 8), "待补充")
    set_cell_text(table.cell(6, 3), "□有  ☑无（公开资料未见海外经历）")
    set_cell_text(table.cell(6, 8), "深圳")
    set_cell_text(table.cell(7, 3), "□是  ☑否（拟按项目需要来苏州）")
    set_cell_text(table.cell(7, 8), "待安排")
    set_cell_text(table.cell(8, 3), "☑全职  □兼职")
    set_cell_text(table.cell(8, 8), "")
    set_cell_text(table.cell(9, 3), "拟作为公司全职商务负责人，负责客户开发、试点合作、销售渠道、合同谈判和产业资源导入。")
    set_cell_text(table.cell(10, 3), "待补充")
    set_cell_text(table.cell(10, 8), "待补充")
    set_cell_text(table.cell(13, 0), "1997.12")
    set_cell_text(table.cell(13, 2), "至今")
    set_cell_text(table.cell(13, 4), "深圳市光彩凯宜电子开发有限公司")
    set_cell_text(table.cell(13, 7), "法定代表人、执行董事&总经理")
    set_cell_text(table.cell(14, 0), "")
    set_cell_text(table.cell(14, 2), "")
    set_cell_text(table.cell(14, 4), "")
    set_cell_text(table.cell(14, 7), "")
    set_cell_text(table.cell(15, 0), "")
    set_cell_text(table.cell(15, 2), "")
    set_cell_text(table.cell(15, 4), "")
    set_cell_text(table.cell(15, 7), "")
    set_cell_text(table.cell(16, 0), "")
    set_cell_text(table.cell(16, 2), "")
    set_cell_text(table.cell(16, 4), "")
    set_cell_text(table.cell(16, 7), "")
    set_cell_text(table.cell(20, 0), "待补充")
    set_cell_text(table.cell(20, 2), "待补充")
    set_cell_text(table.cell(20, 4), "待补充")
    set_cell_text(table.cell(20, 7), "待补充")
    set_cell_text(table.cell(20, 9), "待补充")
    set_cell_text(
        table.cell(26, 0),
        "就业经历和业绩：公开资料显示，马戎先生为深圳市光彩凯宜电子开发有限公司法定代表人、执行董事&总经理。该公司成立于 1997 年，是自动化装备、老化测试装备、精密工装治具、电子电路开发设计制造服务商，公开项目资料列示其主要客户包括华为、H3C、浪潮信息、曙光信息等 ICT 领域企业。马戎先生长期在深圳经营电子开发和自动化测试装备企业，具备面向大型 ICT 客户的商务开拓、客户服务、项目交付和供应链协同经验。\n\n与本项目关联：TLaser 早期商业化需要快速获得光模块、激光芯片、测试装备和 ICT 客户试点。马戎先生拟担任商务负责人，负责苏州及深圳/华南客户资源对接、试点项目商务谈判、销售渠道建设、合同回款和测试装备合作伙伴导入，补足团队在市场开拓和商业闭环方面的能力。"
    )


doc = Document(str(SRC))

summary = (
    "1）团队优势：项目由万振文全职牵头。万振文为中科院博士、国家级人才，曾任厦门大学副教授、丹麦气象研究所资深研究员，近十年聚焦 CAE/工业软件算法创业，完成 EML、MySim 等软件和多项客户合同交付，已取得 3 项发明专利授权。马戎先生拟全职担任商务负责人，公开资料显示其为深圳市光彩凯宜电子开发有限公司法定代表人、执行董事&总经理，长期服务华为、H3C、浪潮信息、曙光信息等 ICT 客户，具备自动化测试装备和大客户商务经验。蒋寻涯教授为复旦大学研究员、博士生导师，曾入选中科院百人计划，作为兼职首席科学顾问负责光电仿真、测试验证技术路线和产业资源。"
    "\n2）目标产品：TLaser 是面向光通讯半导体激光器的数字孪生与在线标定平台，产品包括激光器参数快速仿真、L-I-V 实测曲线反演、批次漂移诊断、良率分析和云端模型管理五个模块。"
    "\n3）核心技术：以准三维激光器物理模型生成训练数据，用 PINN 代理模型实现毫秒级预测，再用实测 L-I-V 数据反向校准内部物性参数，解决传统仿真慢、测试数据与设计模型脱节、批次漂移难诊断的问题。"
    "\n4）客户价值：对光模块、激光芯片和测试装备企业，TLaser 可把器件设计筛选、失效定位和量产一致性分析从“人工经验+离线仿真”变为“数据闭环+快速预测”，减少重复测试和试错迭代，提升研发效率与良率判断能力。"
    "\n5）产业化基础：现已完成模拟器、1500 组训练样本、PINN 模型、标定脚本和中英文控制面板，可开展真实客户数据接入和联合验证。优先服务苏州园区及苏州光通信产业链企业，再复制到长三角和深圳/华南 ICT 客户。"
    "\n6）商业模式与目标：前期以联合验证和模型适配服务收费，中期形成按器件型号/账号/部署方式收费的 SaaS 与私有化部署，第三年形成标准化商业版本。三年目标为新增项目经费 800 万元、主营收入 1330 万元、新增就业 25 人、受理发明专利 5 项、软著 6 项。"
)
replace_para_start(doc, "1）团队优势：项目由万振文全职牵头", summary)

t3 = doc.tables[3]
set_cell_text(
    t3.cell(3, 0),
    "万振文：拟全职，CEO/CTO，负责公司战略、核心算法、产品定义和重点客户交付。其优势是计算物理、工业软件、数值算法和创业交付经验连续，已完成多款 CAE/仿真软件原型和客户合同交付。\n马戎：拟全职，商务负责人/COO，负责苏州、深圳及长三角客户开发，试点合作、销售渠道、合同谈判、回款管理和测试装备合作伙伴导入。公开资料显示，其为深圳市光彩凯宜电子开发有限公司法定代表人、执行董事&总经理，长期经营自动化及老化测试装备、电子电路开发设计制造业务，并服务华为、H3C、浪潮信息、曙光信息等 ICT 客户，具备大客户商务和项目交付经验。\n蒋寻涯：兼职首席科学顾问，负责半导体光电仿真、全波电磁建模、L-I-V 测试验证技术路线、产品技术路线评审和产业客户资源协同。其为复旦大学研究员、博士生导师，曾入选中科院百人计划，并推进 EastWave 国产电磁仿真软件产业化。\n团队分工形成“全职创业与技术负责人 + 全职商务负责人 + 兼职科学顾问”的结构，解决评审关注的市场开拓、产品落地和测试验证责任边界问题。"
)

project_overview = (
    "TLaser 面向光通讯半导体激光器研发与量产测试环节，建设“数字孪生 + 在线标定 + 批次诊断”的工业软件平台。项目负责人万振文为中科院博士、国家级人才，具备科研建模、工业软件、数值算法和创业交付经验；马戎先生拟全职负责商务拓展、客户试点、渠道建设和合同交付闭环；蒋寻涯教授作为兼职首席科学顾问负责光电仿真、L-I-V 测试验证技术路线和产业资源。\n\n"
    "项目背景不是单纯“做一个仿真软件”，而是解决光模块和激光芯片企业在高速产品迭代中的现实痛点：设计仿真慢、测试数据分散、器件批次漂移难解释、失效分析依赖专家经验。目标产品 TLaser 可接入客户 L-I-V 测试曲线，反演内部物性参数，快速预测不同设计参数下的输出功率、效率和空间剖面，并形成批次一致性诊断报告。应用领域包括光通讯激光器芯片设计、光模块研发测试、封装筛选、失效分析、量产质量诊断和测试装备算法嵌入。"
)
replace_para_start(doc, "TLaser 面向光通讯半导体激光器研发与量产测试环节", project_overview)

company_overview = (
    "本项目拟以苏州麦慎数字科技有限公司作为园区产业化主体，注册时间、注册资本和股权结构以工商登记为准。公司定位为光电子 AI 工业软件企业，首个产品为 TLaser 光通讯激光器数字孪生与在线标定平台。当前项目尚未形成规模化营业收入，但已有可运行原型、自动化验证流程、用户手册和演示材料，并具备历史合同交付和发明专利背书。\n\n"
    "落地苏州后，公司第一阶段收入来自联合验证和模型适配服务，第二阶段形成 SaaS 订阅和私有化部署授权，第三阶段扩展到光电子器件数字孪生模型库。公司早期将控制固定成本，以万振文为全职技术/产品核心、马戎为全职商务负责人，配置算法、软件、测试和销售人员；蒋寻涯教授作为兼职首席科学顾问参与关键技术评审、测试验证路线和资源协同。"
)
replace_para_start(doc, "本项目拟以苏州麦慎数字科技有限公司作为园区产业化主体", company_overview)

stage_body = (
    "3-5 年阶段目标围绕产品、客户、团队和知识产权展开。第一年完成苏州公司注册，形成万振文、马戎为核心的全职经营团队，完成 TLaser V1.0 工程版、真实 L-I-V 数据接入、1 家以上苏州客户联合验证，主营收入约 80 万元，申请 1 项发明专利、2 项软著，新增就业 5 人。第二年完成多型号激光器模型库、SaaS 试点版、3 家以上付费客户，主营收入约 350 万元，申请 2 项发明专利、1 项 PCT、2 项软著，新增就业至 12 人。第三年完成商业版、私有化部署和 API 接口，服务不少于 6 类器件型号，主营收入约 900 万元，新增就业至 25 人，累计形成 5 项发明专利受理、2 项 PCT、6 项软著。\n\n"
    "第四至第五年从激光器扩展到硅光调制器、探测器、光模块热电可靠性和封装测试数字孪生，形成光电子器件模型库和行业案例，争取高新技术企业、专精特新培育和新一轮股权融资。"
)
replace_para_start(doc, "3-5 年阶段目标围绕产品、客户、团队和知识产权展开", stage_body)

# Transfer test validation work to Jiang's member table.
for table in doc.tables:
    text = table_text(table)
    if len(table.rows) >= 27 and "蒋寻涯" in text:
        set_cell_text(table.cell(9, 3), "复旦大学在职，拟以首席科学顾问方式参与技术路线评审、光电仿真模型指导、L-I-V 测试验证路线设计和产业资源协同。")
        set_cell_text(
            table.cell(26, 0),
            "就业经历和业绩：蒋寻涯教授为复旦大学研究员、博士生导师，公开资料显示曾在美国 Ames 国家实验室、Dicon Fiberoptics（硅谷）、MIT 电子研究所、中科院上海微系统所、复旦大学等机构从事光子学、光电器件、电磁与光电仿真研究。其研究方向覆盖激光/LED、光子晶体、超材料、特异介质、非线性与无序体系等；公开资料显示其发表 SCI 论文逾百篇，被引用 3000 余次，出版 Springer 学术专著，拥有多项发明专利，并曾入选中科院百人计划、上海杰出引进人才、浦江计划、浙江省千人计划等人才项目。\n\n产业化业绩：蒋寻涯教授创办上海东峻信息科技有限公司，长期推进国产电磁/光电仿真软件 EastWave 产业化。与本项目关联：蒋教授负责 TLaser 光电仿真模型、激光器 L-I-V 测试验证技术路线、测试数据与物理模型一致性评审，以及半导体光电子产业客户资源协同。秦三团原拟承担的测试验证技术职责已调整至蒋教授负责的技术路线和验证体系中。"
        )
    if len(table.rows) >= 27 and "秦三团" in text:
        fill_ma_rong(table)

doc.save(str(OUT))
print(str(OUT))
